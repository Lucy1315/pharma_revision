"""
기존 수작업 보고서와 에이전트 생성 보고서 수치 비교.
사용법: python compare.py
"""
import sys
from pathlib import Path
from docx import Document

BASE = Path(__file__).parent.parent
REFERENCE_DOCX = BASE / "안전관리책임자가 작성한 안전관리에 관한 자료.docx"
GENERATED_DOCX = Path(__file__).parent / "data" / "output" / "generated_report.docx"


def extract_table_data(doc_path: Path) -> dict:
    doc = Document(doc_path)
    tables = doc.tables
    data = {"tables": [], "line_listing_count": 0}

    for i, t in enumerate(tables):
        rows = []
        for row in t.rows:
            rows.append([c.text.strip() for c in row.cells])
        data["tables"].append(rows)

        # Line Listing 탐지: 첫 행에 'KAERS번호' 또는 '번호' 포함
        if rows and any("KAERS" in str(rows[0]) or "번호" in str(rows[0]) for _ in [1]):
            if len(rows) > 1:
                # 헤더 제외 데이터 행 수 = Line Listing 건수
                data["line_listing_count"] = len(rows) - 1

    # 텍스트에서 수치 추출
    full_text = "\n".join(p.text for p in doc.paragraphs)
    data["full_text"] = full_text
    return data


def extract_key_metrics(data: dict, label: str) -> dict:
    metrics = {"label": label}

    # Line Listing 건수 (테이블에서)
    metrics["line_listing_rows"] = data["line_listing_count"]

    # 테이블 수치 추출
    for table in data["tables"]:
        for row in table:
            joined = " ".join(row)
            # 보고건수 테이블: "합계" 행
            if "합계" in row[0] and len(row) >= 4:
                try:
                    metrics["total_reports"] = int(row[-1]) if row[-1].isdigit() else None
                except Exception:
                    pass
            # 중대성 테이블: YES/NO
            if "YES" in joined:
                try:
                    metrics["serious_count"] = int(row[1]) if len(row) > 1 and row[1].isdigit() else None
                except Exception:
                    pass
            if "NO" in joined and "serious_count" in metrics:
                try:
                    metrics["non_serious_count"] = int(row[1]) if len(row) > 1 and row[1].isdigit() else None
                except Exception:
                    pass

    # 텍스트에서 총 건수 추출
    import re
    text = data.get("full_text", "")
    match = re.search(r"총\s*(\d+)\s*건", text)
    if match:
        metrics["total_events_text"] = int(match.group(1))
    match_cases = re.search(r"인원\s*(\d+)\s*명", text)
    if match_cases:
        metrics["total_cases_text"] = int(match_cases.group(1))
    match_male = re.search(r"남성은\s*(\d+)\s*명", text)
    if match_male:
        metrics["male_count"] = int(match_male.group(1))
    match_female = re.search(r"여성은\s*(\d+)\s*명", text)
    if match_female:
        metrics["female_count"] = int(match_female.group(1))

    return metrics


def compare_reports(ref_path: Path, gen_path: Path) -> list[dict]:
    print(f"\n{'='*60}")
    print("보고서 비교 분석")
    print(f"기준(수작업): {ref_path.name}")
    print(f"비교(에이전트): {gen_path.name}")
    print(f"{'='*60}\n")

    ref_data = extract_table_data(ref_path)
    gen_data = extract_table_data(gen_path)

    ref_metrics = extract_key_metrics(ref_data, "수작업")
    gen_metrics = extract_key_metrics(gen_data, "에이전트")

    compare_keys = [
        ("total_events_text", "이상사례 총 건수 (텍스트)"),
        ("total_cases_text", "보고 인원 수 (텍스트)"),
        ("male_count", "남성 수"),
        ("female_count", "여성 수"),
        ("serious_count", "중대한 이상사례 수"),
        ("non_serious_count", "비중대 이상사례 수"),
        ("line_listing_rows", "Line Listing 행 수"),
        ("total_reports", "보고유형 합계"),
    ]

    diffs = []
    print(f"{'항목':<30} {'수작업':>10} {'에이전트':>12} {'일치':>8}")
    print("-" * 64)

    for key, label in compare_keys:
        ref_val = ref_metrics.get(key, "N/A")
        gen_val = gen_metrics.get(key, "N/A")
        match = "✅" if ref_val == gen_val else "❌"
        print(f"{label:<30} {str(ref_val):>10} {str(gen_val):>12} {match:>8}")
        if ref_val != gen_val:
            diffs.append({
                "항목": label, "수작업": ref_val, "에이전트": gen_val,
                "차이": f"{gen_val} - {ref_val}" if isinstance(ref_val, int) and isinstance(gen_val, int) else "불일치",
            })

    print()
    if diffs:
        print(f"❌ 불일치 항목 {len(diffs)}건:")
        for d in diffs:
            print(f"  - {d['항목']}: 수작업={d['수작업']}, 에이전트={d['에이전트']} (차이: {d['차이']})")
    else:
        print("✅ 모든 주요 수치 일치!")

    print(f"\n테이블 수: 수작업={len(ref_data['tables'])}, 에이전트={len(gen_data['tables'])}")
    return diffs


if __name__ == "__main__":
    # 보고서 먼저 생성
    output_path = GENERATED_DOCX
    if not output_path.exists():
        print("에이전트 보고서 생성 중...")
        import subprocess
        result = subprocess.run(
            ["python3", "main.py",
             "--drug-name", "프로테조밉주",
             "--company", "㈜삼양홀딩스",
             "--drug-code", "201506668",
             "--output", str(output_path)],
            capture_output=True, text=True, cwd=Path(__file__).parent
        )
        print(result.stdout)
        if result.returncode != 0:
            print(f"오류: {result.stderr}", file=sys.stderr)
            sys.exit(1)

    if not REFERENCE_DOCX.exists():
        print(f"기준 보고서 없음: {REFERENCE_DOCX}", file=sys.stderr)
        sys.exit(1)

    diffs = compare_reports(REFERENCE_DOCX, output_path)
    sys.exit(0 if not diffs else 1)
