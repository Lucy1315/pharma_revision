import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from .types import ProcessedData
from .aggregator import compute_aggregates


# ── 스타일 헬퍼 ──────────────────────────────────────────

def _yellow_run(para, text: str):
    run = para.add_run(text)
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rPr.append(highlight)
    return run


def _red_run(para, text: str):
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


def _blue_run(para, text: str):
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    return run


def _bold_red_para(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return p


def _add_table_header(table, headers: list[str]):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True


def _add_row(table, values: list):
    row = table.add_row()
    for i, val in enumerate(values):
        row.cells[i].text = str(val) if val is not None else ""


def _para(doc: Document, text: str, bold: bool = False, style: str = "No Spacing") -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold


def _placeholder(doc: Document, label: str):
    p = doc.add_paragraph(style="No Spacing")
    _yellow_run(p, f"[{label}]")


# ── 통계 계산 헬퍼 ──────────────────────────────────────

def _pct(n: int, total: int) -> str:
    if total == 0:
        return "0.0%"
    return f"{n/total*100:.1f}%"


def _compute_stats(data: ProcessedData) -> dict:
    """공유 집계 모듈 호출. Excel과 동일한 수치를 보장."""
    return compute_aggregates(data.df_merged, data.df_line_listing)


# ── 정성 텍스트 생성 ─────────────────────────────────────

def _narrative_report_count(stats: dict, drug_name: str) -> str:
    n = stats["n_cases"]
    q = stats["n_quick"]
    by_type = stats["rpt_cross"]
    parts = []
    for rpt_type in ["시험/연구에서보고", "자발적보고", "기타", "모름"]:
        if rpt_type in by_type.index:
            total = int(by_type.loc[rpt_type].sum())
            if total > 0:
                parts.append(f"{rpt_type}가 {total}건")
    type_str = ", ".join(parts) + "이었다." if parts else ""
    return (
        f"{drug_name}에 관련된 보고건수는 총 {n}건이었으며, "
        f"이 중 신속보고는 {q}건이었다. {type_str}"
    )


def _narrative_demographics(stats: dict) -> str:
    n = stats["n_cases"]
    m, f = stats["male"], stats["female"]
    age_counts = stats["age_counts"]
    age_parts = []
    for age, cnt in age_counts.items():
        if age and age not in ("", "nan"):
            age_parts.append(f"{age} {cnt}명({_pct(cnt, n)})")
    age_str = ", ".join(age_parts[:5]) if age_parts else "정보 없음"
    return (
        f"보고된 인원 {n}명 중 남성은 {m}명({_pct(m, n)})이었으며, "
        f"여성은 {f}명({_pct(f, n)})이었다. "
        f"연령은 {age_str}이었다."
    )


def _narrative_adverse_events(stats: dict) -> str:
    n = stats["n_cases"]
    total_e = stats["n_events"]
    top3 = stats["top3"]
    top_str = ""
    if len(top3) > 0:
        items = []
        for adr, cnt in top3.items():
            items.append(f"{adr} {cnt}건({_pct(cnt, total_e)})")
        top_str = "주요 이상사례는 " + ", ".join(items) + "이었다. "
    serious = stats["n_serious"]
    return (
        f"{n}명의 환자에게 발생한 이상사례는 인과성 평가와 관계없이 총 {total_e}건이었다. "
        f"{top_str}"
        f"확인된 총 {total_e}건의 이상사례 중 중대한 이상사례로 보고된 것은 {serious}건이었다."
    )


# ── (ㄹ)(ㅁ)(ㅂ) 데이터 기반 초안 ─────────────────────────

def _build_approval_comparison(doc: Document, stats: dict, data: ProcessedData):
    """(ㄹ) 허가사항 비교 초안 — 보고된 이상사례 목록을 SOC별로 나열."""
    soc_summary = stats.get("soc_summary")
    n_events = stats.get("n_events", 0)
    drug = data.drug_name or "본 제품"

    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        f"본 분석 기간 중 {drug}과 관련하여 보고된 이상사례는 총 {n_events}건이며, "
        f"기관계대분류(SOC) 기준 분포는 다음과 같다."
    )

    if soc_summary is not None and len(soc_summary) > 0:
        for _, r in soc_summary.iterrows():
            soc = r["SOC_NM"]
            cnt = int(r["건수"])
            pct = f"{cnt / n_events * 100:.1f}%" if n_events else "0.0%"
            _para(doc, f"  • {soc}: {cnt}건 ({pct})")

    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        "상기 이상사례 각각이 현재 허가사항(사용상의 주의사항)에 수재되어 있는지 "
        "여부는 사내 허가사항 문서와의 대조가 필요하다."
    )
    p = doc.add_paragraph(style="No Spacing")
    _blue_run(p, "[검토필요: 상기 SOC/PT 각 항목이 현행 허가사항에 수재되어 있는지 대조]")


def _build_off_label_events(doc: Document, stats: dict, data: ProcessedData):
    """(ㅁ) 허가사항 외 이상사례 초안 — 전체 이상사례를 PT 단위로 나열."""
    pt_summary = stats.get("pt_summary")

    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        "본 분석 기간 중 보고된 이상사례명(PT, Preferred Term) 목록은 아래와 같다. "
        "각 이상사례가 허가사항에 등재되지 않은 신규 이상사례(Off-label)인지 여부는 "
        "상기 (ㄹ) 항목의 허가사항 대조 결과에 따라 판정한다."
    )

    if pt_summary is not None and len(pt_summary) > 0:
        for _, r in pt_summary.iterrows():
            _para(doc, f"  • {r['PT']}: {int(r['건수'])}건")

    p = doc.add_paragraph(style="No Spacing")
    _blue_run(p, "[검토필요: 상기 PT 중 허가사항 미수재 항목을 식별하여 별도 기재]")


def _build_overall_review(doc: Document, stats: dict, data: ProcessedData):
    """(ㅂ) 종합 검토 초안 — 수치 기반 자동 요약문."""
    n = stats.get("n_events", 0)
    n_cases = stats.get("n_cases", 0)
    serious = stats.get("n_serious", 0)
    non_serious = stats.get("n_non_serious", 0)
    serious_pct = f"{serious / n * 100:.1f}%" if n else "0.0%"

    # 인과성 "확실~가능" 집계
    evalt = stats.get("evalt_counts")
    related_terms = {"확실함", "상당히확실함", "가능함"}
    n_related = 0
    if evalt is not None and len(evalt) > 0:
        n_related = int(sum(v for k, v in evalt.items() if str(k) in related_terms))

    drug = data.drug_name or "본 제품"
    start, end = data.analysis_period
    period_str = f"{start} ~ {end}" if start and end else "본 분석 기간"

    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        f"{period_str} 동안 {drug}에 대해 보고된 이상사례는 총 {n}건({n_cases}건의 사례)이며, "
        f"이 중 중대한 이상사례는 {serious}건({serious_pct}), 비중대 이상사례는 {non_serious}건이었다."
    )

    if data.has_assessment:
        p = doc.add_paragraph(style="No Spacing")
        p.add_run(
            f"인과성 평가 결과 '확실함/상당히 확실함/가능함'으로 분류된 이상사례는 {n_related}건이었다."
        )
    else:
        p = doc.add_paragraph(style="No Spacing")
        _blue_run(p, "[검토필요: 인과성 평가 자료(ASSESSMENT.txt) 미제공 — 별도 검토 필요]")

    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        "보고된 이상사례의 SOC별 분포 및 중대성·인과성을 종합 검토한 결과, "
        "(ㄹ)(ㅁ) 항목의 허가사항 대조 결과를 바탕으로 허가사항 개정 필요 여부 및 "
        "신규 안전성 신호(Safety Signal) 발생 여부를 판단한다."
    )
    p = doc.add_paragraph(style="No Spacing")
    _blue_run(p, "[검토필요: 허가사항 개정 필요성 및 안전성 신호 최종 판단]")


# ── 테이블 빌더 ──────────────────────────────────────────

def _build_company_info_table(doc: Document, data: ProcessedData):
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    rows_data = [
        ("회사명", data.company_name, "허가일", data.approval_date),
        ("제품명", data.drug_name, "유효기간", "[이곳에 유효기간을 입력하세요]"),
        ("성분명", data.ingredient_name, "품목기준코드", data.drug_code),
        ("효능효과", "[붙임] 참조", "[붙임] 참조", "[붙임] 참조"),
        ("용법용량", "[붙임] 참조", "[붙임] 참조", "[붙임] 참조"),
    ]
    for i, (c0, c1, c2, c3) in enumerate(rows_data):
        r = table.rows[i]
        r.cells[0].text = c0
        r.cells[1].text = c1
        r.cells[2].text = c2
        r.cells[3].text = c3


def _build_report_type_table(doc: Document, stats: dict):
    rpt_cross = stats["rpt_cross"]
    types = ["시험/연구에서보고", "자발적보고", "기타", "모름"]
    present_types = [t for t in types if t in rpt_cross.index]
    if not present_types:
        present_types = list(rpt_cross.index)

    table = doc.add_table(rows=len(present_types) + 2, cols=4)
    table.style = "Table Grid"
    header_row = table.rows[0]
    for cell, txt in zip(header_row.cells, ["보고유형", "일반보고", "신속보고", "계(단위:건)"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True

    for i, rtype in enumerate(present_types, start=1):
        row = table.rows[i]
        general = int(rpt_cross.loc[rtype, "일반보고"]) if "일반보고" in rpt_cross.columns else 0
        quick = int(rpt_cross.loc[rtype, "신속보고"]) if "신속보고" in rpt_cross.columns else 0
        total = general + quick
        for cell, val in zip(row.cells, [rtype, general, quick, total]):
            cell.text = str(val)

    # 합계 행
    last = table.rows[-1]
    total_general = int(rpt_cross.get("일반보고", pd.Series([0])).sum())
    total_quick = int(rpt_cross.get("신속보고", pd.Series([0])).sum())
    for cell, val in zip(last.cells, ["합계", total_general, total_quick, total_general + total_quick]):
        cell.text = str(val)
        cell.paragraphs[0].runs[0].bold = True


def _build_demographics_table(doc: Document, stats: dict):
    n_cases = stats["n_cases"]
    sex_rows = [
        ("성별", "남성", stats["male"], _pct(stats["male"], n_cases)),
        ("", "여성", stats["female"], _pct(stats["female"], n_cases)),
    ]
    age_counts = stats["age_counts"]
    age_rows = []
    for age, cnt in age_counts.items():
        if age and str(age) not in ("", "nan"):
            age_rows.append(("연령대", age, cnt, _pct(cnt, n_cases)))

    all_rows = sex_rows + age_rows
    table = doc.add_table(rows=len(all_rows) + 2, cols=4)
    table.style = "Table Grid"

    for cell, txt in zip(table.rows[0].cells, ["구분", "구분", "N (단위:명)", "비율(%)"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True

    for i, (grp, sub, n, pct) in enumerate(all_rows, start=1):
        for cell, val in zip(table.rows[i].cells, [grp, sub, n, pct]):
            cell.text = str(val)

    for cell, val in zip(table.rows[-1].cells, ["합계", "", n_cases, "100.0%"]):
        cell.text = str(val)
        cell.paragraphs[0].runs[0].bold = True


def _build_soc_pt_table(doc: Document, stats: dict):
    soc_pt = stats["soc_pt"]
    table = doc.add_table(rows=len(soc_pt) + 2, cols=5)
    table.style = "Table Grid"

    headers = ["기관계대분류", "이상사례명", "발현건수(전체)", "비율", "중대성건수"]
    for cell, h in zip(table.rows[0].cells, headers):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for i, row in soc_pt.iterrows():
        trow = table.rows[i + 1]
        for cell, val in zip(trow.cells, [
            row.get("SOC_NM", ""), row.get("ADR_MEDDRA_KOR_NM", ""),
            row.get("전체건수", 0), row.get("비율", ""), row.get("중대성건수", 0),
        ]):
            cell.text = str(val)

    # 합계 행
    last = table.rows[-1]
    total = int(soc_pt["전체건수"].sum()) if len(soc_pt) > 0 else 0
    serious_total = int(soc_pt["중대성건수"].sum()) if len(soc_pt) > 0 else 0
    for cell, val in zip(last.cells, ["합계", "", total, "100.0%", serious_total]):
        cell.text = str(val)
        cell.paragraphs[0].runs[0].bold = True


def _build_seriousness_table(doc: Document, stats: dict):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    for cell, h in zip(table.rows[0].cells, ["중대성 여부", "이상사례 수 (단위:건)"]):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    n_serious = stats["n_serious"]
    n_non = stats["n_non_serious"]
    total = n_serious + n_non
    for row, vals in zip(table.rows[1:], [
        ("YES", n_serious), ("NO", n_non), ("합계", total)
    ]):
        for cell, val in zip(row.cells, vals):
            cell.text = str(val)


def _build_line_listing_table(doc: Document, ll_df: pd.DataFrame):
    cols = [
        "번호", "KAERS번호", "인과성평가", "이상사례명(한글)",
        "발현일", "종료일", "이상사례경과", "중대성",
    ]
    present = [c for c in cols if c in ll_df.columns]
    n_rows = len(ll_df)
    table = doc.add_table(rows=n_rows + 1, cols=len(present))
    table.style = "Table Grid"

    for cell, h in zip(table.rows[0].cells, present):
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for i, row in ll_df.iterrows():
        trow = table.rows[i + 1]
        for cell, col in zip(trow.cells, present):
            val = str(row.get(col, ""))
            cell.text = val


# ── 메인 빌더 ────────────────────────────────────────────

def build_report(data: ProcessedData) -> bytes:
    doc = Document()
    ll = data.df_line_listing
    stats = _compute_stats(data)
    n_events = stats["n_events"]

    # ── 통계 무결성 검증 ──
    if len(ll) != len(data.df_merged):
        _bold_red_para(
            doc,
            f"⚠️ [통계 무결성 경고] 처리된 이상사례 {len(data.df_merged)}건과 "
            f"Line Listing {len(ll)}행이 불일치합니다. 데이터를 확인하세요.",
        )

    # ── 미확인 코드 경고 ──
    for uc in data.unknown_codes:
        _bold_red_para(doc, f"[미확인코드 경고] {uc['col']} = {uc['code']}")

    # ── 표지 ──
    p = doc.add_paragraph(style="No Spacing")
    run = p.add_run("안전관리책임자가 작성한 안전관리에 관한 자료\n분석·평가 결과 및 안전관리조치에 관한 자료")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 회사 정보 테이블 ──
    _build_company_info_table(doc, data)
    doc.add_paragraph()

    # ── 1. 요약 ──
    _para(doc, "1. 요약", bold=True)
    start, end = data.analysis_period
    summary_text = (
        f"{start}~{end} 동안의 {data.drug_name}에 대한 신속/정기 보고 내역은 모두 0건이었고, "
        f"식품의약품안전처 의약품통합정보시스템으로부터 제공받은 원시자료에서 "
        f"분석 기간 내 이상사례 보고건수는 총 {data.total_cases}건이었다."
    )
    _para(doc, summary_text)
    doc.add_paragraph()

    # ── 2. 상세 ──
    _para(doc, "2. 상세", bold=True)

    # 2.1 약물감시 업무 절차
    _para(doc, "2.1 약물감시 업무 절차")
    p = doc.add_paragraph(style="No Spacing")
    _yellow_run(p, "[이곳에 회사의 약물감시 업무 절차 및 내부 기준서 정보를 입력하세요]")
    doc.add_paragraph()

    # 2.2 갱신 대상 품목 안전관리에 관한 자료
    _para(doc, "2.2 갱신 대상 품목 안전관리에 관한 자료")

    # 가. 신속보고
    _para(doc, "가. 규칙 [별표 4의3] 제7호가목·나목에 따른 신속보고 자료")
    _para(doc, ": 보고 내역 없음")
    doc.add_paragraph()

    # 나. 정기보고
    _para(doc, "나. 규칙 [별표 4의3] 제7호라목·마목에 따른 보고 자료")
    _para(doc, ": 보고 내역 없음")
    doc.add_paragraph()

    # 다. 수집대상정보
    _para(doc, f"다. 수집대상정보(자발적 부작용 보고자료 등) – 제공처: 식품의약품안전처 의약품통합정보시스템")
    _para(doc, f"품목갱신대상 제품인 {data.drug_name} 부작용 보고 자료를 식품의약품안전처 의약품통합정보시스템으로부터 제공받아 분석하였다.")
    doc.add_paragraph()

    # (ㄱ) 보고 건수
    _para(doc, "(ㄱ) 보고 건수")
    _para(doc, _narrative_report_count(stats, data.drug_name))
    doc.add_paragraph()
    _build_report_type_table(doc, stats)
    doc.add_paragraph()

    # (ㄴ) 인구학적 자료
    _para(doc, "(ㄴ) 조사대상자의 인구학적 자료")
    _para(doc, _narrative_demographics(stats))
    doc.add_paragraph()
    _build_demographics_table(doc, stats)
    doc.add_paragraph()

    # (ㄷ) 이상사례 발현 현황
    _para(doc, "(ㄷ) 이상사례 발현 현황")
    _para(doc, _narrative_adverse_events(stats))
    doc.add_paragraph()

    _para(doc, "< 이상사례 발현 현황표 >")
    _build_soc_pt_table(doc, stats)
    doc.add_paragraph()

    _para(doc, "< 중대성 여부 >")
    _build_seriousness_table(doc, stats)
    doc.add_paragraph()

    _para(doc, "< Line Listing >")
    _build_line_listing_table(doc, ll)
    doc.add_paragraph()

    # (ㄹ) 허가사항 비교 — 데이터 기반 자동 초안
    _para(doc, "(ㄹ) 보고된 이상사례와 허가사항의 비교")
    _build_approval_comparison(doc, stats, data)
    doc.add_paragraph()

    # (ㅁ) 허가사항 외 이상사례 — 데이터 기반 자동 초안
    _para(doc, "(ㅁ) 허가사항 외 발생한 이상사례")
    _build_off_label_events(doc, stats, data)
    doc.add_paragraph()

    # (ㅂ) 검토 — 데이터 기반 종합 요약
    _para(doc, "(ㅂ) 검토")
    _build_overall_review(doc, stats, data)
    doc.add_paragraph()

    # ── 경고 로그 ──
    if data.warnings:
        doc.add_paragraph()
        _para(doc, "[처리 중 발생한 경고]", bold=True)
        for w in data.warnings:
            _para(doc, f"  • {w}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
